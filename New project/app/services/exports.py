from __future__ import annotations

from pathlib import Path

from docx import Document

from app.config import Settings
from app.models import BookRecord, ChapterRecord


class ExportService:
    def __init__(self, settings: Settings):
        self.settings = settings
        self.settings.export_dir.mkdir(parents=True, exist_ok=True)

    def export(self, book: BookRecord, chapters: list[ChapterRecord]) -> tuple[str, str]:
        safe_name = "".join(ch for ch in book.title if ch.isalnum() or ch in (" ", "-", "_")).strip()
        stem = safe_name.replace(" ", "_") or book.id
        txt_path = self.settings.export_dir / f"{stem}.txt"
        docx_path = self.settings.export_dir / f"{stem}.docx"

        text = self._build_text(book, chapters)
        txt_path.write_text(text, encoding="utf-8")

        doc = Document()
        doc.add_heading(book.title, level=0)
        for chapter in chapters:
            doc.add_heading(f"Chapter {chapter.chapter_number}: {chapter.chapter_title}", level=1)
            doc.add_paragraph(chapter.content or "")
        doc.save(docx_path)

        return str(Path(txt_path).resolve()), str(Path(docx_path).resolve())

    @staticmethod
    def _build_text(book: BookRecord, chapters: list[ChapterRecord]) -> str:
        parts = [book.title, "", "Outline", book.outline or "", ""]
        for chapter in chapters:
            parts.extend(
                [
                    f"Chapter {chapter.chapter_number}: {chapter.chapter_title}",
                    chapter.content or "",
                    "",
                ]
            )
        return "\n".join(parts).strip() + "\n"
